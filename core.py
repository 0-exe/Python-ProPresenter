import docx
import requests
import re
import os
import google.generativeai as genai
import ast

def detect_songs_with_gemini(text, api_key, model_name):
    """Detects songs in a text using the Gemini API."""
    try:
        if not api_key:
            return "Error: Gemini API key is not set."
        
        genai.configure(api_key=api_key)
        model = genai.GenerativeModel(model_name)
        
        prompt = f"""You are an expert at parsing church service plans. Your task is to extract all song titles from the provided text. Return only a Python list of strings, where each string is a song title.

Here is an example:
---
TEXT:
Anfangslied: Cornerstone (Hillsong)
... some other text ...
Lobpreis-Block
- Mighty to Save
- What a Beautiful Name
... 
Schlusslied: Amazing Grace
---
RESPONSE:
["Cornerstone", "Mighty to Save", "What a Beautiful Name", "Amazing Grace"]
---

Now, perform the same task on the following text. Do not include psalms or other non-song items. If no songs are found, return an empty list `[]`.

TEXT:
{text}
---
RESPONSE:
"""
        
        response = model.generate_content(prompt)

        # Find a Python list-like string in the response text.
        match = re.search(r'\s*(\[.*\])\s*', response.text, re.DOTALL)
        if not match:
            return f"Error: Could not find a list in the Gemini API response. Response text: {response.text}"

        list_string = match.group(1)
        
        try:
            song_list = ast.literal_eval(list_string)
            if not isinstance(song_list, list) or not all(isinstance(s, str) for s in song_list):
                return "Error: Gemini API did not return a valid list of song titles."
            return song_list
        except (ValueError, SyntaxError):
            return f"Error: Failed to parse the list from the Gemini API response: {list_string}"

    except Exception as e:
        return f"Error calling Gemini API: {e}"

def parse_docx(file_path, api_key, model_name):
    """Parses the .docx file to extract songs and psalms from paragraphs and tables."""
    try:
        doc = docx.Document(file_path)
        full_text_parts = []

        # Extract text from paragraphs
        for para in doc.paragraphs:
            full_text_parts.append(para.text)

        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    full_text_parts.append(cell.text)
        
        full_text = "\n".join(full_text_parts)
        
        # Use Gemini to detect songs
        song_titles = detect_songs_with_gemini(full_text, api_key, model_name)

        if isinstance(song_titles, str): # Error occurred
            return song_titles

        items = []
        for title in song_titles:
            items.append({"type": "song", "value": title})

        # Use regex to find unique psalms
        found_psalms = set()
        for para in doc.paragraphs:
            if "Psalm" in para.text:
                match = re.search(r"Psalm (\d+)", para.text)
                if match:
                    psalm_value = f"Psalm {match.group(1)}"
                    if psalm_value not in found_psalms:
                        items.append({"type": "psalm", "value": psalm_value})
                        found_psalms.add(psalm_value)
        
        return items
    except Exception as e:
        return f"Error parsing docx file: {e}"

def get_psalm(psalm_reference, translation):
    """Fetches psalm text from a Bible API."""
    try:
        response = requests.get(f"https://bible-api.com/{psalm_reference}?translation={translation}")
        response.raise_for_status()
        data = response.json()
        return data["text"]
    except requests.exceptions.RequestException as e:
        return f"Error fetching psalm: {e}"

def get_library(propresenter_url):
    """Fetches the ProPresenter library."""
    base_url = propresenter_url.rstrip('/')
    try:
        response = requests.get(f"{base_url}/api/v1/library")
        response.raise_for_status()
        return response.json()
    except requests.exceptions.RequestException as e:
        return f"Error fetching ProPresenter library: {e}"

def find_presentation_in_library(library, name):
    """Finds a presentation in the ProPresenter library."""
    for item in library:
        if item["name"] == name:
            return item
    return None

def create_propresenter_presentation(name, content, propresenter_url):
    """Creates a presentation in ProPresenter."""
    base_url = propresenter_url.rstrip('/')
    try:
        presentation = {
            "name": name,
            "slides": [
                {
                    "text": content
                }
            ]
        }
        response = requests.post(f"{base_url}/api/v1/presentation", json=presentation)
        response.raise_for_status()
        return response.json()
    except requests.exceptions.RequestException as e:
        return f"Error creating ProPresenter presentation: {e}"

def add_to_propresenter_playlist(playlist_id, presentation_id, propresenter_url):
    """Adds a presentation to a playlist in ProPresenter."""
    base_url = propresenter_url.rstrip('/')
    try:
        response = requests.post(f"{base_url}/api/v1/playlist/{playlist_id}/items", json={"id": presentation_id})
        response.raise_for_status()
        return None
    except requests.exceptions.RequestException as e:
        return f"Error adding to ProPresenter playlist: {e}"

def create_propresenter_playlist(playlist_name, items, library, propresenter_url):
    """Creates a playlist in ProPresenter and adds items to it."""
    base_url = propresenter_url.rstrip('/')
    log = []
    try:
        response = requests.post(f"{base_url}/api/v1/playlist", json={"name": playlist_name})
        response.raise_for_status()
        playlist_id = response.json()["id"]
        log.append(f"Successfully created playlist '{playlist_name}' with ID: {playlist_id}")

        for item in items:
            if item["type"] == "song":
                presentation = find_presentation_in_library(library, item["name"])
                if presentation:
                    error = add_to_propresenter_playlist(playlist_id, presentation["id"], base_url)
                    if error:
                        log.append(error)
                    else:
                        log.append(f"Added '{item['name']}' to the playlist.")
                else:
                    log.append(f"Could not find '{item['name']}' in the ProPresenter library.")
            elif item["type"] == "psalm":
                presentation = create_propresenter_presentation(item["name"], item["content"], base_url)
                if presentation and 'id' in presentation:
                    error = add_to_propresenter_playlist(playlist_id, presentation["id"], base_url)
                    if error:
                        log.append(error)
                    else:
                        log.append(f"Added '{item['name']}' to the playlist.")
                else:
                    log.append(f"Failed to create presentation for '{item['name']}'.")
        
        log.append(f"Playlist '{playlist_name}' created successfully.")
        return "\n".join(log)

    except requests.exceptions.RequestException as e:
        error_message = f"Error creating ProPresenter playlist: {e}"
        log.append(error_message)
        return "\n".join(log)
